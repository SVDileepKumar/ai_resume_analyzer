"""PDF and DOCX text extraction service."""

from __future__ import annotations

import asyncio
import io
from fastapi import UploadFile
import pymupdf  # PyMuPDF — high-performance PDF extraction
from docx import Document

# Magic bytes for file type validation (not just extension)
_PDF_MAGIC = b"%PDF"
_DOCX_MAGIC = b"PK\x03\x04"  # DOCX is a ZIP archive

# Minimum chars per page to consider PDF as text-based (not image-only/scanned)
_PDF_MIN_CHARS_PER_PAGE = 50


async def extract_text(file: UploadFile) -> str:
    """Extract text from a PDF or DOCX upload (async-safe).

    Validates file type via magic bytes, not just extension.
    Runs sync PDF/DOCX parsing in a thread to avoid blocking the event loop.
    """
    filename = (file.filename or "").lower()
    file.file.seek(0)
    content = file.file.read()

    if not content:
        raise ValueError(f"{filename}: file is empty.")

    # Validate by magic bytes first, then fallback to extension
    is_pdf = content[:4] == _PDF_MAGIC
    is_docx = content[:4] == _DOCX_MAGIC and filename.endswith(".docx")

    if is_pdf or (not is_docx and filename.endswith(".pdf")):
        if not is_pdf:
            raise ValueError(
                f"{filename}: file extension is .pdf but content is not a valid PDF."
            )
        return await asyncio.to_thread(_extract_pdf, content)

    if is_docx or filename.endswith(".docx"):
        if not is_docx:
            raise ValueError(
                f"{filename}: file extension is .docx but content is not a valid DOCX."
            )
        return await asyncio.to_thread(_extract_docx, content)

    raise ValueError(f"Unsupported file type: {filename}. Use PDF or DOCX.")


def _extract_pdf(content: bytes) -> str:
    """Extract text from raw PDF bytes using PyMuPDF.

    Uses pymupdf's layout-aware text extraction ("blocks" strategy) which
    correctly handles multi-column layouts, reordering text blocks by their
    vertical then horizontal position.  Falls back to raw span extraction if
    blocks yield too little text (e.g. heavily-formatted PDFs).

    Raises ValueError if the PDF appears to be image-only or scanned.
    """
    pages: list[str] = []

    with pymupdf.open(stream=content, filetype="pdf") as doc:
        for page in doc:
            # "blocks" extracts text grouped by visual block, sorted top-to-bottom
            # then left-to-right — ideal for single and multi-column resumes.
            page_text = _extract_page_text(page)
            pages.append(page_text)

    full = "\n".join(pages)
    num_pages = max(1, len(pages))
    total_chars = len(full.replace(" ", "").replace("\n", ""))

    if total_chars < num_pages * _PDF_MIN_CHARS_PER_PAGE:
        raise ValueError(
            "PDF appears to be image-only or scanned; use a text-based PDF or run OCR."
        )
    return full


def _extract_page_text(page: pymupdf.Page) -> str:  # type: ignore[name-defined]
    """Extract text from a single PDF page with layout-aware block ordering.

    Strategy:
    1. Try "blocks" mode — groups text into rectangular blocks sorted by
       (y0, x0), which correctly re-flows multi-column resume layouts.
    2. If blocks yield suspiciously little text vs raw extraction, fall back
       to the raw "text" mode (preserves original reading order from the PDF
       content stream).
    """
    # Extract raw blocks: list of (x0, y0, x1, y1, text, block_no, block_type)
    raw_blocks = page.get_text("blocks", sort=True)  # sort=True: top-to-bottom, L-to-R

    block_lines: list[str] = []
    for block in raw_blocks:
        block_type = block[6]  # 0 = text, 1 = image
        if block_type != 0:
            continue
        text = block[4].strip()
        if text:
            block_lines.append(text)

    blocks_text = "\n".join(block_lines)

    # Sanity-check: if blocks give significantly less than raw mode, fall back
    raw_text = page.get_text("text").strip()
    if len(raw_text) > 0 and len(blocks_text) < len(raw_text) * 0.5:
        return raw_text

    return blocks_text


def _extract_docx(content: bytes) -> str:
    """Extract text from raw DOCX bytes (paragraphs and tables). Tables follow paragraphs."""
    doc = Document(io.BytesIO(content))
    result_paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    result_tables: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text:
                result_tables.append(row_text)
    return "\n".join(result_paras + result_tables)
